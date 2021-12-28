'''
In order to edit word (docx) documents I've used python-docx package.
python-docx parse docx file into a nested tree of the file elements.
in our case we have a few paragraphs and a table, each paragraphs contains
one or more runs which hold the text.
The indexs in the following code have been chosen by trying and guessing
each set of indexs that is in use holds a different editable field.
'''
import docx
from docx.oxml.shared import OxmlElement
import sys

CHECKBOX_PATH = './w:fldChar/w:ffData/w:checkBox'

def generate_docx(args,file_name):
    doc = docx.Document('template.docx')
    doc.paragraphs[0].runs[0].text = args["labName"]
    doc.paragraphs[1].runs[0].text = args["dateCreated"].split('/')[2]
    doc.paragraphs[1].runs[4].text = args["dateCreated"].split('/')[1]
    doc.paragraphs[1].runs[8].text = args["dateCreated"].split('/')[0]
    doc.paragraphs[2].runs[2].text = args["phoneNumber"]
    doc.paragraphs[3].runs[3].text = args["internalNumber"].split('/')[1]
    doc.paragraphs[3].runs[7].text = args["internalNumber"].split('/')[0]
    doc.paragraphs[5].runs[1].text = args["recipient"]
    # This segement handle the checkboxes
    if args["urgency"] == 'normal':
        doc.tables[0].rows[0].cells[0].paragraphs[1].runs[0]._r.xpath(CHECKBOX_PATH)[0].insert(2, OxmlElement('w:checked'))
    if args["urgency"] == 'urgent':
        doc.tables[0].rows[0].cells[0].paragraphs[2].runs[0]._r.xpath(CHECKBOX_PATH)[0].insert(2, OxmlElement('w:checked'))
    if args["urgency"] == 'urgent_arrest':
        doc.tables[0].rows[0].cells[0].paragraphs[3].runs[0]._r.xpath(CHECKBOX_PATH)[0].insert(2, OxmlElement('w:checked'))
    if 'biological' in args["hazards"]:
        doc.tables[0].rows[0].cells[1].paragraphs[1].runs[0]._r.xpath(CHECKBOX_PATH)[0].insert(2, OxmlElement('w:checked'))
    if 'toxic' in args["hazards"]:
        doc.tables[0].rows[0].cells[1].paragraphs[2].runs[0]._r.xpath(CHECKBOX_PATH)[0].insert(2, OxmlElement('w:checked'))
    if 'sharp' in args["hazards"]:
        doc.tables[0].rows[0].cells[1].paragraphs[3].runs[0]._r.xpath(CHECKBOX_PATH)[0].insert(2, OxmlElement('w:checked'))
    if 'normal' in args["exhibits"]:
        doc.tables[0].rows[0].cells[2].paragraphs[1].runs[0]._r.xpath(CHECKBOX_PATH)[0].insert(2, OxmlElement('w:checked'))
    if 'additional' in args["exhibits"]:
        doc.tables[0].rows[0].cells[2].paragraphs[2].runs[0]._r.xpath(CHECKBOX_PATH)[0].insert(2, OxmlElement('w:checked'))
    if 'returning' in args["exhibits"]:
        doc.tables[0].rows[0].cells[2].paragraphs[3].runs[0]._r.xpath(CHECKBOX_PATH)[0].insert(2, OxmlElement('w:checked'))
    doc.tables[0].rows[1].cells[0].paragraphs[0].runs[2].text = args["unit"]
    doc.tables[0].rows[1].cells[0].paragraphs[0].runs[4].text = args["referenceType"]
    doc.tables[0].rows[1].cells[0].paragraphs[0].runs[7].text = args["referenceNumber"]
    doc.tables[0].rows[1].cells[0].paragraphs[1].runs[1].text = args["bagNumber"]
    doc.tables[0].rows[1].cells[0].paragraphs[3].runs[0].text = args["exhibitDescription"]
    doc.tables[0].rows[1].cells[0].paragraphs[5].runs[0].text = args["exhibitsPackaging"]
    doc.tables[0].rows[1].cells[0].paragraphs[7].runs[0].text = args["exhibitsMark"]
    doc.tables[0].rows[1].cells[0].paragraphs[9].runs[0].text = args["eventDescription"]
    doc.tables[0].rows[1].cells[0].paragraphs[11].runs[0].text = args["testingEssense"]
    doc.tables[0].rows[1].cells[0].paragraphs[13].runs[0].text = args["notes"]
    doc.tables[0].rows[1].cells[0].paragraphs[14].runs[1].text = args["senderName"]
    doc.tables[0].rows[1].cells[0].paragraphs[14].runs[3].text = args["senderRank"]
    doc.tables[0].rows[1].cells[0].paragraphs[14].runs[5].text = args["senderSerialNumber"]
    doc.save(f'{file_name}.docx')

def main():
    args = {
        "lab_name":"מעבדת   חבלה   ת\"א",
        "date_created": "23/11/2020",
        "phone_number": "04-5431234",
        "inside_lab": "683/2020",
        "recipient": "אין לי מושג",
        "urgency": "urgent",
        "hazards": ["biological","toxic"],
        "exhibits": ["normal"],
        "unit": "תחנת טבריה",
        "additional_field": "פל\"א",
        "additional_field_info": "8126744",
        "bag_number": "1",
        "exhibit_description": "משהו ממש יפה",
        "exhibits_packaging": "ארוז עם פס",
        "exhibits_mark": "פס אדום",
        "event_description": "מצאנו חפץ חשוד",
        "testing_method": "אין לי מושג למה בדקנו את זה",
        "notes": "שלום",
        "sender_name": "בר שפר",
        "sender_rank": "סמל",
        "sender_serial_number": "05324543"
    }
    generate_docx(args,sys.argv[1])

    

if __name__ == "__main__":
    main() 
'''
In order to edit word (docx) documents I've used python-docx package.
python-docx parse docx file into a nested tree of the file elements.
in our case we have a few paragraphs and a table, each paragraphs contains
one or more runs which hold the text.
The indexs in the following code have been chosen by trying and guessing
each set of indexs that is in use holds a different editable field.
'''

import docx
import io
from docx.oxml.shared import OxmlElement
import os

CHECKBOX_PATH = './w:fldChar/w:ffData/w:checkBox' # the XPath for the checkbox

from docxtpl import DocxTemplate

def generate_transfer_doc(args):
    doc = DocxTemplate("lab_transfer_template.docx")
    context = {
        'lab_name':args['lab_name'],
        'date': args['date'],
         'internal_number':args['internal_number'],
               }
    doc.render(context)
    doc.tables[0].rows[0].cells[0].paragraphs[0].runs[0]._r.xpath(CHECKBOX_PATH)[0]#.insert(2,
                                                                                 #          OxmlElement('w:checked'))
    doc.save("generated_doc.docx")

def generate_docx(args):
    dir_path = os.path.dirname(os.path.realpath(__file__))
    doc = docx.Document(dir_path+'/template.docx') #get  directory of template docx

    doc.paragraphs[0].runs[0].text = 'מעבדת חבלה ' + args["lab_name"]
    doc.paragraphs[1].runs[0].text = '         תאריך: ' + args["date_created"].split('/')[2]
    doc.paragraphs[1].runs[4].text = args["date_created"].split('/')[1]
    doc.paragraphs[1].runs[8].text = args["date_created"].split('/')[0]

    doc.paragraphs[2].runs[2].text = args["phone_number"]
    doc.paragraphs[3].runs[3].text = args["internal_number"]
    doc.paragraphs[5].runs[1].text = args["recipient"]
    # This segement handle the checkboxes
    if args["urgency"] == 'normal':
        doc.tables[0].rows[0].cells[0].paragraphs[1].runs[0]._r.xpath(CHECKBOX_PATH)[0].insert(2,
                                                                                               OxmlElement('w:checked'))
    if args["urgency"] == 'urgent':
        doc.tables[0].rows[0].cells[0].paragraphs[2].runs[0]._r.xpath(CHECKBOX_PATH)[0].insert(2,
                                                                                               OxmlElement('w:checked'))
    if args["urgency"] == 'urgent_arrest':
        doc.tables[0].rows[0].cells[0].paragraphs[3].runs[0]._r.xpath(CHECKBOX_PATH)[0].insert(2,
                                                                                               OxmlElement('w:checked'))
    if 'biological' in args["hazards"]:
        doc.tables[0].rows[0].cells[1].paragraphs[1].runs[0]._r.xpath(CHECKBOX_PATH)[0].insert(2,
                                                                                               OxmlElement('w:checked'))
    if 'toxic' in args["hazards"]:
        doc.tables[0].rows[0].cells[1].paragraphs[2].runs[0]._r.xpath(CHECKBOX_PATH)[0].insert(2,
                                                                                               OxmlElement('w:checked'))
    if 'sharp' in args["hazards"]:
        doc.tables[0].rows[0].cells[1].paragraphs[3].runs[0]._r.xpath(CHECKBOX_PATH)[0].insert(2,
                                                                                                OxmlElement('w:checked'))
    if args["exhibits"] == 'normal':
        doc.tables[0].rows[0].cells[2].paragraphs[1].runs[0]._r.xpath(CHECKBOX_PATH)[0].insert(2,
                                                                                               OxmlElement('w:checked'))
    if args["exhibits"] == 'additional':
        doc.tables[0].rows[0].cells[2].paragraphs[2].runs[0]._r.xpath(CHECKBOX_PATH)[0].insert(2,
                                                                                               OxmlElement('w:checked'))
    if args["exhibits"] == 'returning':
        doc.tables[0].rows[0].cells[2].paragraphs[3].runs[0]._r.xpath(CHECKBOX_PATH)[0].insert(2,
                                                                                               OxmlElement('w:checked'))
    doc.tables[0].rows[1].cells[0].paragraphs[0].runs[2].text = "TEST"
    doc.tables[0].rows[1].cells[0].paragraphs[2].runs[0].text = args["exhibit_description"]

    doc.tables[0].rows[1].cells[0].paragraphs[6].runs[0].text = args["testing_essence"]
    doc.tables[0].rows[1].cells[0].paragraphs[8].runs[0].text = args["notes"]
    doc.tables[0].rows[1].cells[0].paragraphs[9].runs[1].text = args["name"]
    doc.tables[0].rows[1].cells[0].paragraphs[9].runs[3].text = args["rank"]
    doc.tables[0].rows[1].cells[0].paragraphs[9].runs[5].text = args["person_id"]
    # save the file to a buffer that will be sent to the frontend.
    buffer = io.BytesIO() 
    doc.save(buffer)
    #doc.save("generated_doc.docx")
    buffer.seek(0) 
    return buffer


def main():
    args = {
        'lab_name':"LAB",#
        'phone_number':"PHONE",#
        'internal_number': "IN",#
        'recipient': "RECEPEINT",#
        'urgency': "urgent",#
        'hazards': "sharp",#
        'exhibits': "exhibits",#
        "testing_essence": "testing_essence",  #
        "notes": "notes",  #
        "name": "name",  #
        "rank": "rank",  #
        "person_id": "person_id",#
        "reference_type": "reference_type",
        "reference_number": "פלאא",
        "exhibit_description": "1. "+'what_sampled1'+" ממוצג מס' "+'exhibit_number1'+' בדוח התפיסה הוכנסו לשקית צלף שסומנה "'+'packaging1'+'" והוכנסה לשקית מאובטחת לשימוש חד פעמי שמספרה '+ str(1) +'\n'+"2. "+'what_sampled2'+" ממוצג מס' "+'exhibit_id2'+' בדוח התפיסה הוכנסו לשקית צלף שסומנה "'+'packaging2'+'" והוכנסה לשקית מאובטחת לשימוש חד פעמי שמספרה '+ str(2) +'\n'"3. "+'what_sampled3'+" ממוצג מס' "+'exhibit_id3'+' בדוח התפיסה הוכנסו לשקית צלף שסומנה "'+'packaging3'+'" והוכנסה לשקית מאובטחת לשימוש חד פעמי שמספרה '+ str(3) +'\n',
        "event_description": "event_description",
        'date_created': "08/08/2005"



    }
    #generate_docx(args)


if __name__ == "__main__":
    main()